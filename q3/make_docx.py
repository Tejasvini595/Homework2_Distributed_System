import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import re
import os

doc = docx.Document()

# Page margins
sections = doc.sections
for section in sections:
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

# Read ANALYSIS.md
with open('ANALYSIS.md', 'r') as f:
    lines = f.readlines()

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)

in_table = False
table_lines = []
in_code = False
code_lines = []

def process_table(t_lines):
    rows_data = []
    for line in t_lines:
        line = line.strip()
        if not line or line.startswith('|---') or line.startswith('| :---'):
            continue
        parts = [p.strip() for p in line.split('|')[1:-1]]
        if parts:
            rows_data.append(parts)
    if not rows_data:
        return
    
    num_cols = max(len(r) for r in rows_data)
    table = doc.add_table(rows=len(rows_data), cols=num_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    for r_idx, row in enumerate(rows_data):
        for c_idx, val in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            clean_val = re.sub(r'[\*\$\`]', '', val)
            cell.text = clean_val
            p = cell.paragraphs[0]
            if p.runs:
                p.runs[0].font.size = Pt(9.5)
                if r_idx == 0:
                    set_cell_background(cell, '1F4E78')
                    p.runs[0].font.color.rgb = RGBColor(255, 255, 255)
                    p.runs[0].font.bold = True
                else:
                    if r_idx % 2 == 1:
                        set_cell_background(cell, 'F2F4F7')
                    else:
                        set_cell_background(cell, 'FFFFFF')
    doc.add_paragraph() # spacer

for line in lines:
    raw_line = line.rstrip('\n')
    line_str = raw_line.strip()

    # Code block handling
    if line_str.startswith('```'):
        if in_code:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run('\n'.join(code_lines))
            run.font.name = 'Courier New'
            run.font.size = Pt(9.5)
            code_lines = []
            in_code = False
        else:
            in_code = True
            code_lines = []
        continue

    if in_code:
        code_lines.append(raw_line)
        continue

    # Table handling
    if line_str.startswith('|'):
        in_table = True
        table_lines.append(line_str)
        continue
    else:
        if in_table:
            process_table(table_lines)
            table_lines = []
            in_table = False

    if not line_str or line_str == '---':
        continue

    # Image handling
    img_match = re.match(r'!\[(.*?)\]\((.*?)\)', line_str)
    if img_match:
        caption, img_path = img_match.groups()
        if os.path.exists(img_path):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run()
            run.add_picture(img_path, width=Inches(5.8))
            
            cp = doc.add_paragraph()
            cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cp.paragraph_format.space_after = Pt(12)
            crun = cp.add_run(f'Figure: {caption}')
            crun.font.italic = True
            crun.font.size = Pt(9.5)
            crun.font.color.rgb = RGBColor(100, 100, 100)
        continue

    # Headings
    if line_str.startswith('# '):
        p = doc.add_heading(level=1)
        run = p.add_run(line_str[2:].replace('**', '').replace('$', ''))
        run.font.size = Pt(18)
        run.font.color.rgb = RGBColor(31, 78, 120)
        continue
    elif line_str.startswith('## '):
        p = doc.add_heading(level=2)
        run = p.add_run(line_str[3:].replace('**', '').replace('$', ''))
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(46, 117, 182)
        continue
    elif line_str.startswith('### '):
        p = doc.add_heading(level=3)
        run = p.add_run(line_str[4:].replace('**', '').replace('$', ''))
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(89, 89, 89)
        continue

    # List items
    if line_str.startswith('- ') or line_str.startswith('* '):
        p = doc.add_paragraph(style='List Bullet')
        text = line_str[2:]
        parts = re.split(r'(\*\*.*?\*\*)', text)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                r = p.add_run(part[2:-2].replace('$', ''))
                r.bold = True
            else:
                p.add_run(part.replace('$', ''))
        continue

    # Paragraph text
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    parts = re.split(r'(\*\*.*?\*\*)', line_str)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            r = p.add_run(part[2:-2].replace('$', ''))
            r.bold = True
        else:
            p.add_run(part.replace('$', ''))

if in_table:
    process_table(table_lines)

doc.save('ANALYSIS.docx')
print('Successfully generated ANALYSIS.docx!')
